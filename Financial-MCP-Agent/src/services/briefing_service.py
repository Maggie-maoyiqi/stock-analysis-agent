"""早晚简报生成服务。"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from .profile_service import load_profile
from .recommendation_service import (
    build_position_reviews,
    build_watchlist_reviews,
    discover_recommendations,
    fetch_market_context,
)

PROJECT_ROOT = Path(__file__).resolve().parents[2]
BRIEF_DIR = PROJECT_ROOT / "reports" / "briefs"


def _markdown_table(headers: List[str], rows: List[List[Any]]) -> str:
    table = "| " + " | ".join(headers) + " |\n"
    table += "|" + "|".join([" --- " for _ in headers]) + "|\n"
    for row in rows:
        table += "| " + " | ".join(str(cell) for cell in row) + " |\n"
    return table


def _save_docx(report: Dict[str, Any], docx_path: Path):
    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.styles["Normal"].font.name = "PingFang SC"
    document.styles["Normal"].font.size = Pt(10.5)

    document.add_heading(report["title"], level=0)
    document.add_paragraph(f"生成时间：{report['generated_at']}")
    document.add_paragraph(f"风险偏好：{report['risk_preference']}  |  主战场：A股")

    document.add_heading("一、海外与宏观速览", level=1)
    macro_table = document.add_table(rows=1, cols=3)
    macro_table.style = "Table Grid"
    for idx, header in enumerate(["指标", "最新值", "涨跌幅"]):
        macro_table.rows[0].cells[idx].text = header
    for item in report["market_context"].get("items", []):
        row = macro_table.add_row().cells
        row[0].text = item["label"]
        row[1].text = str(item["latest"])
        row[2].text = f"{item['change_pct']}%"

    document.add_heading("二、自选股观察", level=1)
    watch_table = document.add_table(rows=1, cols=6)
    watch_table.style = "Table Grid"
    for idx, header in enumerate(["代码", "名称", "5日涨幅", "20日波动率", "综合分", "建议"]):
        watch_table.rows[0].cells[idx].text = header
    for item in report["watchlist_reviews"]:
        row = watch_table.add_row().cells
        row[0].text = item["stock_code"]
        row[1].text = item["stock_name"]
        row[2].text = f"{item.get('recent_change_pct_5d', 0)}%"
        row[3].text = f"{item.get('volatility_20d_pct', 0)}%"
        row[4].text = str(item.get("overall_score", ""))
        row[5].text = item.get("action", "")

    document.add_heading("三、持仓体检", level=1)
    if report["position_reviews"]:
        position_table = document.add_table(rows=1, cols=6)
        position_table.style = "Table Grid"
        for idx, header in enumerate(["代码", "名称", "盈亏%", "回撤%", "动作", "备注"]):
            position_table.rows[0].cells[idx].text = header
        for item in report["position_reviews"]:
            row = position_table.add_row().cells
            row[0].text = item["stock_code"]
            row[1].text = item["stock_name"]
            row[2].text = f"{item.get('pnl_pct', 0)}%"
            row[3].text = f"{item.get('drawdown_from_high_pct', 0)}%"
            row[4].text = f"{item.get('risk_tag', '')} {item.get('action', '')}"
            row[5].text = "；".join(item.get("reasons", [])[:2])
    else:
        document.add_paragraph("当前没有录入持仓。")

    document.add_heading("四、推荐买入 5 只", level=1)
    for index, item in enumerate(report["recommendations"], start=1):
        document.add_heading(f"{index}. {item['stock_name']} ({item['stock_code']})", level=2)
        document.add_paragraph(
            f"当前价 {item['current_price']}，综合分 {item['overall_score']}，"
            f"预测方向 {item['forecast_direction']}，建议关注区间 {item['buy_zone']}。"
        )
        document.add_paragraph(item["analysis"])
        document.add_paragraph(f"风险点：{item['risk_point']}")

    document.add_heading("五、风险提示", level=1)
    document.add_paragraph("本简报为决策辅助，不构成投资建议。请严格执行止损纪律，并结合个人账户情况独立决策。")
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def _render_markdown(report: Dict[str, Any]) -> str:
    market_rows = [[item["label"], item["latest"], f"{item['change_pct']}%"] for item in report["market_context"].get("items", [])]
    watch_rows = [
        [
            item["stock_code"],
            item["stock_name"],
            f"{item.get('recent_change_pct_5d', 0)}%",
            f"{item.get('volatility_20d_pct', 0)}%",
            item.get("overall_score", ""),
            item.get("action", ""),
        ]
        for item in report["watchlist_reviews"]
    ]
    position_rows = [
        [
            item["stock_code"],
            item["stock_name"],
            f"{item.get('pnl_pct', 0)}%",
            f"{item.get('drawdown_from_high_pct', 0)}%",
            f"{item.get('risk_tag', '')} {item.get('action', '')}",
            "；".join(item.get("reasons", [])[:2]),
        ]
        for item in report["position_reviews"]
    ]

    recommendation_blocks = []
    for index, item in enumerate(report["recommendations"], start=1):
        recommendation_blocks.append(
            "\n".join(
                [
                    f"### {index}. {item['stock_name']} ({item['stock_code']})",
                    f"- 当前价：{item['current_price']}",
                    f"- 综合分：{item['overall_score']}",
                    f"- 预测方向：{item['forecast_direction']} ({item['forecast_return_pct']}%)",
                    f"- 建议买入区间：{item['buy_zone']}",
                    f"- 推荐理由：{item['analysis']}",
                    f"- 风险点：{item['risk_point']}",
                ]
            )
        )

    sections = [
        f"# {report['title']}",
        f"- 生成时间：{report['generated_at']}",
        f"- 风险偏好：{report['risk_preference']}",
        "## 一、海外与宏观速览",
        _markdown_table(["指标", "最新值", "涨跌幅"], market_rows) if market_rows else "暂无数据",
        "## 二、自选股观察",
        _markdown_table(["代码", "名称", "5日涨幅", "20日波动率", "综合分", "建议"], watch_rows) if watch_rows else "暂无自选股",
        "## 三、持仓体检",
        _markdown_table(["代码", "名称", "盈亏%", "回撤%", "动作", "备注"], position_rows) if position_rows else "当前暂无持仓",
        "## 四、推荐买入 5 只",
        "\n\n".join(recommendation_blocks) if recommendation_blocks else "当前没有满足稳健条件的推荐股票。",
        "## 五、风险提示",
        "本简报仅供决策辅助，不构成投资建议。",
    ]
    return "\n\n".join(sections)


async def generate_daily_brief(session: str = "morning") -> Dict[str, Any]:
    """生成早报或晚报，并导出 Markdown/Word。"""
    profile = load_profile()
    market_context = await fetch_market_context(session)
    watchlist_reviews = await build_watchlist_reviews(profile)
    position_reviews = await build_position_reviews(profile)
    recommendations = await discover_recommendations(profile)

    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    title_prefix = "早间盘前简报" if session == "morning" else "晚间复盘简报"
    title = f"{title_prefix} - {datetime.now().strftime('%Y-%m-%d')}"
    report = {
        "title": title,
        "generated_at": generated_at,
        "session": session,
        "risk_preference": profile.get("risk_preference", "稳健"),
        "market_context": market_context,
        "watchlist_reviews": watchlist_reviews,
        "position_reviews": position_reviews,
        "recommendations": recommendations,
    }

    BRIEF_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    markdown_path = BRIEF_DIR / f"brief_{session}_{timestamp}.md"
    docx_path = BRIEF_DIR / f"brief_{session}_{timestamp}.docx"

    markdown = _render_markdown(report)
    markdown_path.write_text(markdown, encoding="utf-8")
    docx_output = str(docx_path)
    try:
        _save_docx(report, docx_path)
    except Exception:
        docx_output = ""

    report.update(
        {
            "markdown": markdown,
            "markdown_file": str(markdown_path),
            "docx_file": docx_output,
        }
    )
    return report
